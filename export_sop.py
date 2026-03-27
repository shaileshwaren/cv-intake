"""
Export SOP.md → SOP.docx

Usage:  python export_sop.py
        python export_sop.py --input OTHER.md --output OTHER.docx

Requires: pip install -r requirements-docs.txt
"""

import argparse
import os
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SCRIPT_DIR = Path(__file__).parent
DEFAULT_INPUT = SCRIPT_DIR / "SOP.md"
DEFAULT_OUTPUT = SCRIPT_DIR / "SOP.docx"
SCREENSHOT_DIR = SCRIPT_DIR / "screenshots"

FONT_NAME = "Calibri"
FONT_SIZE_BODY = Pt(11)
FONT_SIZE_H1 = Pt(22)
FONT_SIZE_H2 = Pt(16)
FONT_SIZE_H3 = Pt(13)
FONT_SIZE_H4 = Pt(11)

TABLE_HEADER_COLOR = "D9E2F3"
TABLE_BORDER_COLOR = "B4C6E7"


def set_font(run, size=FONT_SIZE_BODY, bold=False, italic=False, color=None):
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f"</w:tblBorders>"
    )
    tbl_pr.append(borders)


def add_formatted_text(paragraph, text, base_size=FONT_SIZE_BODY):
    """Parse inline markdown (bold, italic, code, bold-italic) and add runs."""
    pattern = re.compile(
        r"(\*\*\*(.+?)\*\*\*)"   # bold italic
        r"|(\*\*(.+?)\*\*)"      # bold
        r"|(\*(.+?)\*)"          # italic
        r"|(`(.+?)`)"            # inline code
    )
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
            set_font(run, size=base_size)

        if m.group(2):  # bold italic
            run = paragraph.add_run(m.group(2))
            set_font(run, size=base_size, bold=True, italic=True)
        elif m.group(4):  # bold
            run = paragraph.add_run(m.group(4))
            set_font(run, size=base_size, bold=True)
        elif m.group(6):  # italic
            run = paragraph.add_run(m.group(6))
            set_font(run, size=base_size, italic=True)
        elif m.group(8):  # code
            run = paragraph.add_run(m.group(8))
            run.font.name = "Consolas"
            run.font.size = Pt(base_size.pt - 1) if hasattr(base_size, "pt") else Pt(10)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        pos = m.end()

    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_font(run, size=base_size)


def resolve_image_path(md_path: str) -> Path | None:
    """Turn a markdown image path into an absolute path, if the file exists."""
    p = SCREENSHOT_DIR.parent / md_path
    if p.exists():
        return p
    p = SCREENSHOT_DIR / Path(md_path).name
    if p.exists():
        return p
    return None


def parse_table_block(lines: list[str]) -> list[list[str]]:
    """Parse consecutive markdown table lines into a list of rows (list of cell strings)."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            break
        cells = [c.strip() for c in line.strip("|").split("|")]
        if all(set(c) <= {"-", ":", " "} for c in cells):
            continue
        rows.append(cells)
    return rows


def add_table_to_doc(doc, rows: list[list[str]]):
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j >= n_cols:
                break
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            is_header = i == 0
            add_formatted_text(p, cell_text, base_size=FONT_SIZE_BODY)
            for run in p.runs:
                if is_header:
                    run.font.bold = True
            if is_header:
                set_cell_shading(cell, TABLE_HEADER_COLOR)

    doc.add_paragraph()


def build_docx(md_path: Path, docx_path: Path):
    md_text = md_path.read_text(encoding="utf-8")
    lines = md_text.split("\n")

    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE_BODY
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    for level, size in [(1, FONT_SIZE_H1), (2, FONT_SIZE_H2), (3, FONT_SIZE_H3), (4, FONT_SIZE_H4)]:
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = FONT_NAME
        hs.font.size = size
        hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        hs.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Horizontal rule
        if stripped in ("---", "***", "___"):
            doc.add_paragraph()
            i += 1
            continue

        # Headings
        heading_match = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            p = doc.add_heading(level=level)
            add_formatted_text(p, text, base_size=[FONT_SIZE_H1, FONT_SIZE_H2, FONT_SIZE_H3, FONT_SIZE_H4][level - 1])
            i += 1
            continue

        # Image (with optional caption on next line)
        img_match = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if img_match:
            alt_text = img_match.group(1)
            img_rel = img_match.group(2)
            img_path = resolve_image_path(img_rel)
            if img_path:
                doc.add_picture(str(img_path), width=Inches(5.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p = doc.add_paragraph()
                run = p.add_run(f"[Image placeholder: {alt_text}]")
                set_font(run, italic=True, color=RGBColor(0x88, 0x88, 0x88))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Check for italic caption on next line
            if i + 1 < len(lines):
                caption_match = re.match(r"^\s*\*\[(.+)\]\*\s*$", lines[i + 1].strip())
                if caption_match:
                    cap_p = doc.add_paragraph()
                    cap_run = cap_p.add_run(caption_match.group(1))
                    set_font(cap_run, size=Pt(9), italic=True, color=RGBColor(0x66, 0x66, 0x66))
                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    i += 1
            i += 1
            continue

        # Italic-only line (screenshot caption)
        caption_match = re.match(r"^\*\[(.+)\]\*$", stripped)
        if caption_match:
            p = doc.add_paragraph()
            run = p.add_run(caption_match.group(1))
            set_font(run, size=Pt(9), italic=True, color=RGBColor(0x66, 0x66, 0x66))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Table block
        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            table_lines = []
            j = i
            while j < len(lines) and lines[j].strip().startswith("|"):
                table_lines.append(lines[j])
                j += 1
            rows = parse_table_block(table_lines)
            add_table_to_doc(doc, rows)
            i = j
            continue

        # Checkbox list item
        checkbox_match = re.match(r"^- \[( |x)\]\s+(.*)", stripped)
        if checkbox_match:
            checked = checkbox_match.group(1) == "x"
            text = checkbox_match.group(2)
            p = doc.add_paragraph(style="List Bullet")
            prefix = "\u2611 " if checked else "\u2610 "
            run = p.add_run(prefix)
            set_font(run, size=FONT_SIZE_BODY)
            add_formatted_text(p, text)
            i += 1
            continue

        # Numbered list item (e.g. "1. text" or "**1. text**")
        num_match = re.match(r"^\*?\*?(\d+)\.\s+(.*?)\*?\*?\s*$", stripped)
        if num_match and not stripped.startswith("|"):
            text = stripped
            # Clean leading ** and trailing ** if the whole line is bold
            if text.startswith("**") and text.endswith("**"):
                text = text[2:-2]
            p = doc.add_paragraph(style="List Number")
            # Remove the number prefix since the style provides numbering
            inner = re.sub(r"^\d+\.\s+", "", text)
            add_formatted_text(p, inner)
            i += 1
            continue

        # Bullet list item
        bullet_match = re.match(r"^[-*]\s+(.*)", stripped)
        if bullet_match:
            text = bullet_match.group(1)
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_text(p, text)
            i += 1
            continue

        # Empty line
        if not stripped:
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        add_formatted_text(p, stripped)
        i += 1

    doc.save(str(docx_path))
    print(f"Exported: {docx_path}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Export SOP.md to SOP.docx")
    parser.add_argument("--input", type=Path, default=DEFAULT_INPUT, help="Markdown source")
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT, help="Word output")
    args = parser.parse_args()
    build_docx(args.input, args.output)
